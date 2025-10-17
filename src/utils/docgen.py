from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_horizontal_line(doc):
    paragraph = doc.add_paragraph()
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert(0, pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def generate_assignment_docx(
    questions,
    filename,
    subject_name,
    assignment_number,
    student_name,
    student_class,
    student_div,
    student_rollno,
    student_batch,
):
    doc = Document()

    # Centered header
    hdr = doc.add_paragraph(f"{subject_name} Assignment {assignment_number}")
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Centered student info
    info = doc.add_paragraph(f"Name: {student_name}\nClass: {student_class}\nDiv: {student_div}\nRoll No: {student_rollno}\nBatch: {student_batch}")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Horizontal line
    add_horizontal_line(doc)

    for q in questions:
        doc.add_paragraph(f"Question {q['number']}", style="Heading 2")
        doc.add_paragraph(q['text'])
        doc.add_paragraph("Code:", style="Heading 3")
        code_par = doc.add_paragraph()
        code_run = code_par.add_run(q['code'])
        code_run.font.name = 'Consolas'
        doc.add_paragraph("Output:", style="Heading 3")
        doc.add_paragraph(q['output'] or "[No Output]")
        add_horizontal_line(doc)

    doc.save(filename)
    return filename
